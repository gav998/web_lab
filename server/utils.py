import datetime
import sqlite3
import os
import subprocess
import json
from sys import platform
from collections import Counter

from docx import Document
from htmldocx import HtmlToDocx
from docx.shared import Mm

from io import BytesIO 


CURR_PATH = os.path.abspath(os.getcwd())
DB_PATH = CURR_PATH + "/data.db"
#DB_PATH = CURR_PATH + "\data.db"

 
def tbl_new_uuid(UUID, TEST, NAME, NUM, LETTER): 
    test = generate_test(TEST) 
    sql = f'''
    INSERT 
      INTO tests(UUID, TEST, NAME, NUM, LETTER, TIME_START, {', '.join(list(test.keys()))}) 
    VALUES (?, ?, ?, ?, ?, ?, {', '.join(['?' for i in list(test.keys())])});    
    ''' 
    values = [UUID, TEST, NAME.upper(), NUM, LETTER, datetime.datetime.today()]  
    values.extend(test.values())
    values = tuple(values)
    with sqlite3.connect(DB_PATH) as db:
        db.execute(sql, values)
    return {
                "UUID": UUID,
           }
        
def tbl_get_test(UUID, TASK):
    sql = '''
    SELECT * FROM tests WHERE UUID = (?);
    '''   
    with sqlite3.connect(DB_PATH) as db:
        column = db.execute(sql, (UUID,))
        column_names = [desc[0] for desc in column.description]
        column = column.fetchone()
        test = dict(zip(column_names, column))
    return {
                    "UUID": test["UUID"],
                    "TEST": test["TEST"],
                    "NAME": test["NAME"],
                    "NUM": test["NUM"],
                    "LETTER": test["LETTER"],
                    "TIME_START": test["TIME_START"],
                    "COUNT": test["COUNT"],
                    "LOCK": test["LOCK"],
                    
                    "TASK_TEXT": test["T_"+TASK+"_TEXT"],
                    "ANSW": test["T_"+TASK+"_ANSW"],
                    "ANSW_TIME": test["T_"+TASK+"_ANSW_TIME"],
                    
                }

"""
# на raspberry не поддерживается RETURNING в sqlite
def tbl_get_full_test_and_lock(UUID):
    sql = f'''
    UPDATE tests 
       SET 
        LOCK = 1
     WHERE UUID = (?)
     RETURNING *;
    '''
    with sqlite3.connect(DB_PATH) as db:
        column = db.execute(sql, (UUID,))
        column_names = [desc[0] for desc in column.description]
        column = column.fetchone()
        test = dict(zip(column_names, column))
        return test
"""
def tbl_get_full_test_and_lock(UUID):
    sql1 = f'''
    UPDATE tests 
       SET 
        LOCK = 1
     WHERE UUID = (?);
    '''
    sql2 = f'''SELECT * FROM tests WHERE UUID = (?);'''
    with sqlite3.connect(DB_PATH) as db:
        db.execute(sql1, (UUID,)).fetchone()
        column = db.execute(sql2, (UUID,))
        column_names = [desc[0] for desc in column.description]
        column = column.fetchone()
        test = dict(zip(column_names, column))
    return test

     
def tbl_set_answ(UUID, task, answ):
    sql = f'''
    UPDATE tests 
       SET 
        T_{int(task)}_ANSW = (?),
        T_{int(task)}_ANSW_TIME = (?)
     WHERE UUID = (?) and LOCK = 0;
    '''
    with sqlite3.connect(DB_PATH) as db:
        answ_time = datetime.datetime.today()
        db.execute(sql, (answ, answ_time, UUID,))
    return {
                    "ANSW_TIME": str(answ_time),
                }


def generate_test(TEST):
    if not TEST in os.listdir(path='./tests'):
        return {}
        
    if "win" in platform:
        result = subprocess.run(['python', f'./tests/{TEST}'],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE)
    else:
        result = subprocess.run(['python3', f'./tests/{TEST}'],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE)
    
    if "win" in platform:
        d = json.loads(result.stdout.decode('cp1251'))
    else:
        d = json.loads(result.stdout.decode('utf-8'))

    return d
    
    
def tbl_get_db():
    sql = f'''
    SELECT * FROM tests
    ORDER BY NAME, TIME_START DESC;
    '''
    with sqlite3.connect(DB_PATH) as db:
        column = db.execute(sql)
        column_names = [desc[0] for desc in column.description]
        column = column.fetchall()
        
    data = [dict(zip(column_names, column[i])) for i in range(len(column))]
    
    print(data[0])
    
    return data
    
def tbl_get_headers():
    sql_0 = 'SELECT substring(TIME_START, 1, 10) FROM tests ORDER BY TIME_START DESC;'
    sql_1 = 'SELECT NUM FROM tests ORDER BY NUM;'
    sql_2 = 'SELECT ("NUM" || "--" || "LETTER") FROM tests ORDER BY NUM, LETTER;'
    sql_3 = 'SELECT ("NUM" || "--" || "LETTER" || "--" || "TEST") FROM tests ORDER BY NUM, LETTER, TEST;'
    sql_4 = 'SELECT TEST FROM tests ORDER BY TEST;'
    sql_5 = 'SELECT (substring(TIME_START, 1, 10) || "--" || "NUM" || "--" || "LETTER" || "--" || "TEST") FROM tests ORDER BY TIME_START DESC, NUM, LETTER, TEST;'
    with sqlite3.connect(DB_PATH) as db: 
        db.row_factory = lambda cursor, row: row[0]
        d0 = dict(Counter(db.execute(sql_0).fetchall()))
        d1 = dict(Counter(db.execute(sql_1).fetchall()))
        d2 = dict(Counter(db.execute(sql_2).fetchall()))
        d3 = dict(Counter(db.execute(sql_3).fetchall()))
        d4 = dict(Counter(db.execute(sql_4).fetchall()))
        d5 = dict(Counter(db.execute(sql_5).fetchall()))
    return {
        0:d0,
        1:d1,
        2:d2,
        3:d3,
        4:d4,
        5:d5,
    }

def tbl_get_results(TIME_START, NUM, LETTER, TEST):
    sql = f'''
    SELECT *
      FROM tests
     WHERE TIME_START LIKE (?) || "%" AND 
           NUM LIKE (?) AND 
           LETTER LIKE (?) AND 
           TEST LIKE (?) 
     ORDER BY NUM,
              LETTER,
              NAME,
              TEST,
              TIME_START DESC;
    '''
    with sqlite3.connect(DB_PATH) as db:
        column = db.execute(sql, (TIME_START, NUM, LETTER, TEST,))
        column_names = [desc[0] for desc in column.description]
        column = column.fetchall()  
    data = [dict(zip(column_names, column[i])) for i in range(len(column))]
    return data
    
    
def get_40_tests(TEST):
    # Преднастройка документа и парсера html
    document = Document()
    new_parser = HtmlToDocx()
    section = document.sections[0]
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    # генерация 40 вариантов
    sres = ''
    s = ''
    var = 1  
    while var <= 40:
        document.add_paragraph(f'Вариант {var}\n')
        sres += f"\nОтветы. Вариант {var}\n"
        test1 = generate_test(TEST)
        for i in range(1, 1+test1['COUNT']):
            html = f'{i}. '+test1[f'T_{i}_TEXT']
            new_parser.add_html_to_document(html, document)
            sres += f' {i}.'+test1[f'T_{i}_ANSW_CORRECT']
        document.add_page_break()
        var += 1
    document.add_paragraph(sres)
    
    # сохраняем в поток
    target_stream = BytesIO()
    document.save(target_stream)
    print('готово')
    return target_stream.getvalue()
